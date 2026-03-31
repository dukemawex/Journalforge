# apps/api/src/scripts/build_docx.py
"""
build_docx.py — JournalForge DOCX assembly script.

Usage:
    python3 build_docx.py <actions_json_path> <output_docx_path>

Reads a JSON array of assembly action objects produced by the AI assembler stage
and constructs a python-docx Document, writing the result to <output_docx_path>.

NOTE on display equations:
    This script renders LaTeX strings as plain text inside a centred paragraph
    with the equation label. Production deployments should replace the
    add_display_equation handler with a python-docx-oxml OMML conversion function
    (e.g. using latex2mathml + a LaTeX-to-OMML converter) so that equations render
    natively in Word's equation editor.
"""

import sys
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ALIGNMENT_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def set_paragraph_font(paragraph, font_name="Times New Roman", font_size=12):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def add_run_formatted(paragraph, text, font_name="Times New Roman", font_size=12,
                      bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return run


def set_line_spacing(paragraph, spacing=2.0):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn('w:line'), str(int(spacing * 240)))
    spacing_el.set(qn('w:lineRule'), 'auto')
    spacing_el.set(qn('w:before'), '0')
    spacing_el.set(qn('w:after'), '72')  # 6pt * 12 = 72 twips
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing_el)


def add_continuous_line_numbering(section):
    """Add continuous line numbering to a section via OOXML."""
    sectPr = section._sectPr
    lnNumEl = OxmlElement('w:lnNumType')
    lnNumEl.set(qn('w:countBy'), '1')
    lnNumEl.set(qn('w:restart'), 'continuous')
    lnNumEl.set(qn('w:distance'), '720')
    existing = sectPr.find(qn('w:lnNumType'))
    if existing is not None:
        sectPr.remove(existing)
    sectPr.append(lnNumEl)


def handle_set_page_layout(doc, action):
    section = doc.sections[0]
    margins_cm = float(action.get("margins_cm", 2.54))
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4
    section.left_margin = Cm(margins_cm)
    section.right_margin = Cm(margins_cm)
    section.top_margin = Cm(margins_cm)
    section.bottom_margin = Cm(margins_cm)
    # Store line_spacing for use in subsequent paragraphs
    doc._jf_line_spacing = float(action.get("line_spacing", 2.0))


def handle_add_line_numbering(doc, action):
    section = doc.sections[0]
    add_continuous_line_numbering(section)


def handle_add_title(doc, action):
    para = doc.add_paragraph()
    alignment = ALIGNMENT_MAP.get(action.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    para.alignment = alignment
    add_run_formatted(
        para,
        action["content"],
        font_size=int(action.get("font_size", 12)),
        bold=bool(action.get("bold", True)),
    )
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_author_line(doc, action):
    para = doc.add_paragraph()
    alignment = ALIGNMENT_MAP.get(action.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    para.alignment = alignment
    add_run_formatted(para, action["content"], font_size=int(action.get("font_size", 12)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_affiliation(doc, action):
    para = doc.add_paragraph()
    alignment = ALIGNMENT_MAP.get(action.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    para.alignment = alignment
    add_run_formatted(para, action["content"], font_size=int(action.get("font_size", 10)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_corresponding(doc, action):
    para = doc.add_paragraph()
    alignment = ALIGNMENT_MAP.get(action.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    para.alignment = alignment
    add_run_formatted(para, action["content"], font_size=int(action.get("font_size", 10)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_section_heading(doc, action):
    para = doc.add_paragraph()
    level = int(action.get("level", 1))
    bold = bool(action.get("bold", True))
    italic = bool(action.get("italic", False))
    font_size = int(action.get("font_size", 12))

    if level == 1:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT

    alignment_override = ALIGNMENT_MAP.get(action.get("alignment", ""), None)
    if alignment_override is not None:
        alignment = alignment_override

    para.alignment = alignment
    add_run_formatted(para, action["content"], font_size=font_size, bold=bold, italic=italic)
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_metadata_label(doc, action):
    para = doc.add_paragraph()
    add_run_formatted(
        para,
        action["content"],
        font_size=int(action.get("font_size", 12)),
        bold=bool(action.get("bold", True)),
    )
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_paragraph(doc, action):
    para = doc.add_paragraph()
    alignment = ALIGNMENT_MAP.get(action.get("alignment", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)
    para.alignment = alignment
    add_run_formatted(para, action["content"], font_size=int(action.get("font_size", 12)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_display_equation(doc, action):
    """
    Renders the LaTeX string as plain text centred on its own line.
    Production deployments should swap this for an OMML conversion using
    latex2mathml + a LaTeX-to-OMML converter so that equations render natively
    in Word's equation editor.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    latex_text = action.get("latex", "")
    label = action.get("label", "")
    display_text = f"{latex_text}    {label}" if label else latex_text
    add_run_formatted(para, display_text, font_size=int(action.get("font_size", 12)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_figure_placeholder(doc, action):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = f"[Figure {action['number']} — {action['caption']}]"
    add_run_formatted(
        para, caption,
        font_size=int(action.get("font_size", 10)),
        italic=True,
    )
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_table_placeholder(doc, action):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption = f"[Table {action['number']} — {action['caption']}]"
    add_run_formatted(
        para, caption,
        font_size=int(action.get("font_size", 10)),
        italic=True,
    )
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


def handle_add_section_break(doc, action):
    """Add a page break before the References section."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def handle_add_reference_entry(doc, action):
    """Add a reference entry with hanging indent."""
    para = doc.add_paragraph()
    hanging_cm = float(action.get("hanging_indent_cm", 0.635))
    pPr = para._p.get_or_add_pPr()
    ind_el = OxmlElement('w:ind')
    # hanging indent: left = hanging, hanging = hanging (twips = cm * 567)
    twips = int(hanging_cm * 567)
    ind_el.set(qn('w:left'), str(twips))
    ind_el.set(qn('w:hanging'), str(twips))
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(ind_el)
    add_run_formatted(para, action["content"], font_size=int(action.get("font_size", 12)))
    set_line_spacing(para, getattr(doc, '_jf_line_spacing', 2.0))


HANDLERS = {
    "set_page_layout": handle_set_page_layout,
    "add_line_numbering": handle_add_line_numbering,
    "add_title": handle_add_title,
    "add_author_line": handle_add_author_line,
    "add_affiliation": handle_add_affiliation,
    "add_corresponding": handle_add_corresponding,
    "add_section_heading": handle_add_section_heading,
    "add_metadata_label": handle_add_metadata_label,
    "add_paragraph": handle_add_paragraph,
    "add_display_equation": handle_add_display_equation,
    "add_figure_placeholder": handle_add_figure_placeholder,
    "add_table_placeholder": handle_add_table_placeholder,
    "add_section_break": handle_add_section_break,
    "add_reference_entry": handle_add_reference_entry,
}


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 build_docx.py <actions_json_path> <output_docx_path>", file=sys.stderr)
        sys.exit(1)

    actions_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(actions_path):
        print(f"Actions file not found: {actions_path}", file=sys.stderr)
        sys.exit(1)

    with open(actions_path, "r", encoding="utf-8") as f:
        actions = json.load(f)

    if not isinstance(actions, list):
        print("Actions JSON must be an array.", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    # Remove default empty paragraph
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    doc._jf_line_spacing = 2.0  # default; overridden by set_page_layout

    for i, action in enumerate(actions):
        action_type = action.get("type")
        if not action_type:
            print(f"Warning: action at index {i} has no 'type' field, skipping.", file=sys.stderr)
            continue
        handler = HANDLERS.get(action_type)
        if handler is None:
            print(f"Warning: unknown action type '{action_type}' at index {i}, skipping.", file=sys.stderr)
            continue
        try:
            handler(doc, action)
        except Exception as e:
            print(f"Error processing action '{action_type}' at index {i}: {e}", file=sys.stderr)
            sys.exit(1)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Document saved to {output_path}")


if __name__ == "__main__":
    main()
